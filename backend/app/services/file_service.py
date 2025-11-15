"""
Advanced File Processing Service
Handles secure file uploads, validation, text extraction, and storage management.
"""

import os
import uuid
import magic
import hashlib
import asyncio
from pathlib import Path
from typing import Optional, List, Dict, Any, BinaryIO
from datetime import datetime
import logging

from fastapi import UploadFile, HTTPException, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import aiofiles
from minio import Minio
from minio.error import S3Error
from pypdf import PdfReader
from docx import Document
import spacy
from io import BytesIO

from app.core.config import settings
from app.models.resume import Resume
from app.models.user import User

logger = logging.getLogger(__name__)

# Load spaCy model for text processing
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logger.warning("spaCy English model not found. Text processing will be limited.")
    nlp = None


class FileValidationError(HTTPException):
    """Custom file validation error."""
    
    def __init__(self, detail: str):
        super().__init__(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=detail
        )


class FileStorageError(HTTPException):
    """Custom file storage error."""
    
    def __init__(self, detail: str = "File storage error"):
        super().__init__(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=detail
        )


class FileService:
    """Advanced file processing and storage service."""
    
    def __init__(self):
        """Initialize storage client based on environment."""
        try:
            self.s3_client = Minio(
                settings.storage_endpoint,
                access_key=settings.storage_access_key,
                secret_key=settings.storage_secret_key,
                secure=settings.storage_secure
            )
            self._validate_storage_access()
            self.storage_available = True
        except Exception as e:
            logger.warning(f"Storage service initialization warning: {e}")
            logger.warning("File uploads will fail until storage is properly configured")
            self.s3_client = None
            self.storage_available = False
    
    def _validate_storage_access(self):
        """Validate access to the storage bucket."""
        if not self.s3_client:
            return
            
        try:
            provider = "R2" if settings.is_production else "MinIO"
            if not settings.is_production:
                # For MinIO, create bucket if it doesn't exist
                if not self.s3_client.bucket_exists(settings.storage_bucket_name):
                    self.s3_client.make_bucket(settings.storage_bucket_name)
                    logger.info(f"Created MinIO bucket: {settings.storage_bucket_name}")
            else:
                # For R2, just check if we can list objects
                self.s3_client.list_objects(settings.storage_bucket_name, prefix="", recursive=False)
            
            logger.info(f"Successfully connected to {provider} bucket: {settings.storage_bucket_name}")
            self.storage_available = True
        except Exception as e:
            logger.warning(f"Storage access warning ({provider}): {e}")
            logger.warning(f"Storage service may not be available. File uploads will fail until {provider} is configured properly")
            self.storage_available = False
            # Don't raise - allow app to start
    
    @staticmethod
    def validate_file_type(file: UploadFile) -> tuple[bool, str]:
        """
        Validate file type using magic numbers and extension.
        
        Args:
            file: Uploaded file
            
        Returns:
            tuple: (is_valid, error_message)
        """
        # Check file extension
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in ['.pdf', '.docx']:
            return False, f"Unsupported file type: {file_extension}. Only PDF and DOCX files are allowed."
        
        # Read first few bytes for magic number validation
        file.file.seek(0)
        file_header = file.file.read(1024)
        file.file.seek(0)
        
        try:
            file_type = magic.from_buffer(file_header, mime=True)
            
            valid_mime_types = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
            
            expected_mime = valid_mime_types.get(file_extension)
            if file_type != expected_mime:
                return False, f"File content doesn't match extension. Expected {expected_mime}, got {file_type}"
            
        except Exception as e:
            logger.warning(f"Magic number validation failed: {e}")
            # Fallback to extension-based validation
        
        return True, ""
    
    @staticmethod
    def validate_file_size(file: UploadFile) -> tuple[bool, str]:
        """
        Validate file size.
        
        Args:
            file: Uploaded file
            
        Returns:
            tuple: (is_valid, error_message)
        """
        # Get file size
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning
        
        if file_size > settings.max_file_size:
            size_mb = file_size / (1024 * 1024)
            max_size_mb = settings.max_file_size / (1024 * 1024)
            return False, f"File too large: {size_mb:.1f}MB. Maximum allowed: {max_size_mb:.1f}MB"
        
        if file_size == 0:
            return False, "File is empty"
        
        return True, ""
    
    @staticmethod
    async def scan_file_for_malware(file_content: bytes) -> tuple[bool, str]:
        """
        Basic malware scanning using heuristics.
        
        Args:
            file_content: File content as bytes
            
        Returns:
            tuple: (is_safe, warning_message)
        """
        # Check for suspicious patterns
        suspicious_patterns = [
            b'<script',
            b'javascript:',
            b'eval(',
            b'exec(',
            b'system(',
            b'shell_exec',
            b'passthru'
        ]
        
        content_lower = file_content.lower()
        for pattern in suspicious_patterns:
            if pattern in content_lower:
                return False, f"Suspicious content detected: {pattern.decode('utf-8', errors='ignore')}"
        
        # Check file size vs content ratio (basic heuristic)
        if len(file_content) > 50 * 1024 * 1024:  # 50MB
            return False, "File suspiciously large for document type"
        
        return True, ""
    
    async def extract_text_from_pdf(self, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file content
            
        Returns:
            tuple: (extracted_text, metadata)
        """
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            text_content = []
            metadata = {
                "num_pages": len(pdf_reader.pages),
                "extraction_method": "pypdf2",
                "has_images": False,
                "encryption_status": "encrypted" if pdf_reader.is_encrypted else "not_encrypted"
            }
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    continue
            
            extracted_text = "\n\n".join(text_content)
            
            if not extracted_text.strip():
                raise ValueError("No text could be extracted from PDF")
            
            # Basic content validation
            if len(extracted_text) < 50:
                logger.warning("Extracted text is very short, might be image-based PDF")
            
            return extracted_text, metadata
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise FileValidationError(f"Failed to extract text from PDF: {str(e)}")
    
    async def extract_text_from_docx(self, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content
            
        Returns:
            tuple: (extracted_text, metadata)
        """
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = []
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "extraction_method": "python-docx",
                "has_headers_footers": False
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # Check for headers/footers
            for section in doc.sections:
                if section.header.paragraphs:
                    metadata["has_headers_footers"] = True
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(f"[HEADER: {paragraph.text}]")
                
                if section.footer.paragraphs:
                    metadata["has_headers_footers"] = True
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(f"[FOOTER: {paragraph.text}]")
            
            extracted_text = "\n\n".join(text_content)
            
            if not extracted_text.strip():
                raise ValueError("No text could be extracted from DOCX")
            
            return extracted_text, metadata
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise FileValidationError(f"Failed to extract text from DOCX: {str(e)}")
    
    def preprocess_text(self, text: str) -> Dict[str, Any]:
        """
        Preprocess extracted text using NLP.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Dictionary with processed text and analysis
        """
        processed_data = {
            "raw_text": text,
            "cleaned_text": "",
            "word_count": 0,
            "sentence_count": 0,
            "key_sections": [],
            "entities": [],
            "skills_detected": [],
            "language": "en"
        }
        
        try:
            # Basic text cleaning
            lines = text.split('\n')
            cleaned_lines = []
            
            for line in lines:
                # Remove excessive whitespace
                line = ' '.join(line.split())
                if line and len(line) > 2:  # Skip very short lines
                    cleaned_lines.append(line)
            
            cleaned_text = '\n'.join(cleaned_lines)
            processed_data["cleaned_text"] = cleaned_text
            processed_data["word_count"] = len(cleaned_text.split())
            
            # Use spaCy if available
            if nlp and cleaned_text:
                doc = nlp(cleaned_text[:1000000])  # Limit text size for processing
                
                # Count sentences
                processed_data["sentence_count"] = len(list(doc.sents))
                
                # Extract named entities
                entities = []
                for ent in doc.ents:
                    if ent.label_ in ["PERSON", "ORG", "GPE", "DATE"]:
                        entities.append({
                            "text": ent.text,
                            "label": ent.label_,
                            "confidence": ent._.get("confidence", 0.8)
                        })
                
                processed_data["entities"] = entities[:20]  # Limit entities
                
                # Detect potential skills (basic keyword matching)
                skill_keywords = {
                    "python", "java", "javascript", "react", "node.js", "sql", "aws", "docker",
                    "kubernetes", "git", "agile", "scrum", "machine learning", "ai", "data science",
                    "project management", "leadership", "communication", "teamwork"
                }
                
                text_lower = cleaned_text.lower()
                detected_skills = []
                for skill in skill_keywords:
                    if skill in text_lower:
                        detected_skills.append(skill)
                
                processed_data["skills_detected"] = detected_skills
                
                # Identify key sections
                section_patterns = [
                    "experience", "education", "skills", "projects", "certifications",
                    "summary", "objective", "achievements", "awards"
                ]
                
                key_sections = []
                for pattern in section_patterns:
                    if pattern in text_lower:
                        key_sections.append(pattern)
                
                processed_data["key_sections"] = key_sections
            
        except Exception as e:
            logger.error(f"Text preprocessing error: {e}")
            # Return basic processed data even if NLP fails
            processed_data["cleaned_text"] = text
            processed_data["word_count"] = len(text.split())
        
        return processed_data
    
    async def upload_file_to_storage(self, file_content: bytes, filename: str, content_type: str) -> str:
        """
        Upload file to MinIO storage or save locally as fallback.
        
        Args:
            file_content: File content as bytes
            filename: Filename for storage
            content_type: MIME type of the file
            
        Returns:
            Storage path/key
        """
        # Try MinIO/R2 storage first
        if self.storage_available and self.s3_client is not None:
            try:
                file_stream = BytesIO(file_content)
                
                self.s3_client.put_object(
                    bucket_name=settings.storage_bucket_name,
                    object_name=filename,
                    data=file_stream,
                    length=len(file_content),
                    content_type=content_type
                )
                
                logger.info(f"File uploaded to storage: {filename}")
                return filename
                
            except S3Error as e:
                logger.error(f"MinIO upload error: {e}")
                raise FileUploadError(f"Failed to upload file to storage: {str(e)}")
        
        # Fallback to local file storage
        logger.warning("Storage service unavailable, using local file storage")
        try:
            # Create uploads directory if it doesn't exist
            upload_dir = Path("uploads")
            upload_dir.mkdir(exist_ok=True)
            
            # Create user subdirectory
            user_dir = upload_dir / Path(filename).parent
            user_dir.mkdir(parents=True, exist_ok=True)
            
            # Save file locally
            file_path = upload_dir / filename
            file_path.write_bytes(file_content)
            
            logger.info(f"File saved locally: {file_path}")
            return filename
            
        except Exception as e:
            logger.error(f"Local file save error: {e}")
            raise FileUploadError(f"Failed to save file: {str(e)}")
            raise FileStorageError(f"Failed to upload file to storage: {str(e)}")
    
    async def delete_file_from_storage(self, filename: str):
        """
        Delete file from MinIO storage.
        
        Args:
            filename: Filename to delete
        """
        try:
            self.s3_client.remove_object(settings.storage_bucket_name, filename)
            logger.info(f"File deleted from storage: {filename}")
        except S3Error as e:
            provider = "R2" if settings.is_production else "MinIO"
            logger.error(f"{provider} delete error: {e}")
            raise FileStorageError(f"Failed to delete file from storage: {str(e)}")
    
    async def process_resume_file(
        self, 
        file: UploadFile, 
        user: User, 
        db: AsyncSession
    ) -> Resume:
        """
        Complete resume file processing pipeline.
        
        Args:
            file: Uploaded file
            user: User who uploaded the file
            db: Database session
            
        Returns:
            Created Resume object
        """
        # Validation phase
        logger.info(f"Processing resume file: {file.filename} for user: {user.email}")
        
        # Validate file type
        is_valid_type, type_error = self.validate_file_type(file)
        if not is_valid_type:
            raise FileValidationError(type_error)
        
        # Validate file size
        is_valid_size, size_error = self.validate_file_size(file)
        if not is_valid_size:
            raise FileValidationError(size_error)
        
        # Read file content
        file_content = await file.read()
        file.file.seek(0)  # Reset file pointer
        
        # Malware scan
        is_safe, scan_warning = await self.scan_file_for_malware(file_content)
        if not is_safe:
            raise FileValidationError(f"Security scan failed: {scan_warning}")
        
        # Generate unique filename
        file_extension = Path(file.filename).suffix.lower()
        unique_filename = f"{user.id}/{uuid.uuid4()}{file_extension}"
        
        try:
            # Extract text based on file type
            if file_extension == '.pdf':
                extracted_text, extraction_metadata = await self.extract_text_from_pdf(file_content)
            elif file_extension == '.docx':
                extracted_text, extraction_metadata = await self.extract_text_from_docx(file_content)
            else:
                raise FileValidationError("Unsupported file type")
            
            # Preprocess text
            processed_data = self.preprocess_text(extracted_text)
            
            # Upload file to storage
            await self.upload_file_to_storage(file_content, unique_filename, file.content_type)
            
            # Create resume record in database
            resume = Resume(
                user_id=user.id,
                filename=unique_filename,
                original_filename=file.filename,
                file_path=unique_filename,
                file_size=len(file_content),
                mime_type=file.content_type,
                extracted_text=processed_data["cleaned_text"],
                is_processed=True,
                processing_status="completed",
                processed_at=datetime.utcnow()
            )
            
            db.add(resume)
            await db.commit()
            await db.refresh(resume)
            
            logger.info(f"Resume processed successfully: {resume.id}")
            return resume
            
        except Exception as e:
            # Cleanup on error
            try:
                await self.delete_file_from_storage(unique_filename)
            except:
                pass
            
            logger.error(f"Resume processing failed: {e}")
            raise FileValidationError(f"File processing failed: {str(e)}")
    
    async def get_resume_file_url(self, resume: Resume, expires_in_hours: int = 24) -> str:
        """
        Generate a presigned URL for downloading a resume file.
        
        Args:
            resume: Resume object
            expires_in_hours: URL expiration time in hours
            
        Returns:
            Presigned download URL
        """
        try:
            from datetime import timedelta
            
            url = self.s3_client.presigned_get_object(
                bucket_name=settings.storage_bucket_name,
                object_name=resume.file_path,
                expires=timedelta(hours=expires_in_hours)
            )
            
            return url
            
        except S3Error as e:
            logger.error(f"Error generating presigned URL: {e}")
            raise FileStorageError("Failed to generate download URL")


# Global file service instance
file_service = FileService()